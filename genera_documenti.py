"""
Generatore Documenti v7.1 — Comune di Sanremo, Servizio Demanio Marittimo
Template ricchi con dati strutturati dal DB. Niente più placeholder generici.
Integrazione AI opzionale per le premesse.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def fmtE(valore):
    if valore is None or valore == 0:
        return "€ ___"
    s = f"{valore:,.2f}"
    return "€ " + s.replace(",", "X").replace(".", ",").replace("X", ".")


def fmtD(data_str):
    if not data_str:
        return "___________"
    try:
        return datetime.strptime(data_str, "%Y-%m-%d").strftime("%d/%m/%Y")
    except ValueError:
        return data_str


def _o(val, fallback="___"):
    return str(val).strip() if val and str(val).strip() else fallback


class DocBuilder:
    def __init__(self):
        self.doc = Document()
        sec = self.doc.sections[0]
        sec.page_width = Cm(21); sec.page_height = Cm(29.7)
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2)
        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"; style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.15

    def _run(self, para, text, bold=False, italic=False, size=11):
        run = para.add_run(text)
        run.font.name = "Times New Roman"; run.font.size = Pt(size)
        run.bold = bold; run.italic = italic
        r = run._element
        rPr = r.find(qn("w:rPr"))
        if rPr is None:
            rPr = r.makeelement(qn("w:rPr"), {}); r.insert(0, rPr)
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = rPr.makeelement(qn("w:rFonts"), {}); rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        return run

    def bold(self, text):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._run(p, text, bold=True); return p

    def normal(self, text):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._run(p, text); return p

    def mixed(self, *parts):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for flag, text in parts:
            self._run(p, text, bold=('B' in flag.upper()), italic=('I' in flag.upper()))
        return p

    def center(self, text, bold=False):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(p, text, bold=bold); return p

    def right(self, text, bold=False):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._run(p, text, bold=bold); return p

    def empty(self):
        self.doc.add_paragraph()

    def table(self, data, headers=None, widths=None):
        """Crea una tabella. data è una lista di liste/tuple."""
        rows = len(data)
        cols = len(headers) if headers else (len(data[0]) if data else 0)
        if cols == 0: return
        
        t = self.doc.add_table(rows=(1 if headers else 0) + rows, cols=cols)
        t.style = "Table Grid"
        
        # Header
        if headers:
            for i, h in enumerate(headers):
                cell = t.rows[0].cells[i]
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._run(p, str(h), bold=True, size=10)
        
        # Dati
        start_row = 1 if headers else 0
        for i, row_data in enumerate(data):
            for j, val in enumerate(row_data):
                cell = t.rows[start_row + i].cells[j]
                p = cell.paragraphs[0]
                # Allineamento: se è un numero (contiene € o è float/int) a destra
                if isinstance(val, (int, float)) or (isinstance(val, str) and "€" in val):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                self._run(p, str(val) if val is not None else "", size=10)
        
        if widths:
            for i, w in enumerate(widths):
                for cell in t.columns[i].cells:
                    cell.width = Cm(w)
        
        self.empty()
        return t

    def save(self, path):
        self.doc.save(path); return path


def _intestazione(b, d):
    import database as db
    b.center("COMUNE DI SANREMO", bold=True)
    b.center("Provincia di Imperia")
    b.empty()
    b.bold(db.get_impostazione("settore_nome", "Settore Sviluppo Economico, Ambientale e Floricoltura"))
    b.bold(db.get_impostazione("servizio_nome", "Servizio Demanio Marittimo"))
    b.empty()


def _visti_sanremo(b, include_tracciabilita=True):
    b.bold("VISTI")
    b.normal("– il vigente Statuto Comunale;")
    b.normal("– il vigente Regolamento Comunale di Contabilità;")
    b.normal("– il vigente Regolamento Comunale sui Controlli Interni;")
    b.normal("– il Decreto Legislativo 18 agosto 2000, n. 267 (TUEL);")
    b.normal("– il Decreto Legislativo 31 marzo 2023, n. 36 (Codice dei contratti pubblici);")
    if include_tracciabilita:
        b.normal("– la Legge 13 agosto 2010, n. 136 in materia di tracciabilità dei flussi finanziari;")
    b.normal("– il Decreto Legislativo 14 marzo 2013, n. 33 (Amministrazione Trasparente);")

def _attestazione_copertura(b, d, coperture):
    b.empty(); b.empty()
    b.center("ATTESTAZIONE DI COPERTURA FINANZIARIA", bold=True); b.empty()
    b.normal("Si attesta, ai sensi dell'art. 183, comma 7, del D.Lgs. 267/2000, la regolarità contabile "
             "e la copertura finanziaria della spesa di cui alla presente determinazione.")
    b.empty()
    
    if coperture:
        for c in coperture:
            rows = [
                ["Esercizio Finanziario", _o(c.get("anno_bilancio"), _o(d.get("esercizio")))],
                ["Missione / Programma", f"{c.get('missione','_')} / {c.get('programma','_')}"],
                ["Titolo / Macroaggregato", f"{c.get('titolo','_')} / {c.get('macroaggregato','_')}"],
                ["Capitolo", f"{c.get('capitolo','_')} - {c.get('nome_capitolo','')}"],
                ["Importo", fmtE(c.get("importo", 0))]
            ]
            b.table(rows, widths=[6, 12])
    else:
        b.normal("Dati finanziari non disponibili nel sistema."); b.empty()

    b.empty()
    b.normal("Visto di regolarità contabile favorevole (art. 147-bis D.Lgs. 267/2000).")
    b.empty(); b.empty()
    b.right("Il Responsabile del Servizio Finanziario", bold=True)
    b.right("(Firmato digitalmente)")

def _firma(b, d):
    b.empty(); b.empty()
    b.right(_o(d.get("dirigente"), "Il Dirigente"), bold=True)
    b.right(_o(d.get("qualifica_dirigente"), ""))

def _premesse_ai_o_manuali(b, d, ai_premesse, default_text):
    if ai_premesse:
        for para in ai_premesse.split("\n\n"):
            if para.strip(): b.normal(para.strip()); b.empty()
    else:
        # Se non c'è AI, usa le premesse salvate nella determina o nell'affidamento
        has_manual = False
        for i in range(1, 4):
            val = d.get(f"premessa_{i}")
            if val and val.strip():
                b.normal(val.strip()); b.empty()
                has_manual = True
        
        if not has_manual:
            b.normal(default_text); b.empty()

def _fmt_cop(c):
    return f"Cap. {_o(c.get('capitolo'))} / Art. {_o(c.get('articolo','0'))} - " \
           f"Esercizio {c.get('anno_bilancio','_')} - Importo {fmtE(c.get('importo',0))}"


def _coperture_testo(b, d, copertures):
    if not copertures:
        b.normal(f"la spesa trova copertura al Cap. {_o(d.get('capitolo_bilancio', d.get('aff_capitolo')))} "
                 f"del Bilancio {_o(d.get('esercizio'))};")
    else:
        b.normal("la spesa trova copertura come segue:")
        for c in copertures:
            b.normal(f"– {_fmt_cop(c)};")


def _visti_comuni(b, include_tracciabilita=True):
    return _visti_sanremo(b, include_tracciabilita)


def genera_determina(d, path, coperture=None, qe=None, ai_premesse=None):
    tipo = d.get("tipo_determina")
    if tipo == "Impegno/Accertamento":
        return _gen_impegno_accertamento(d, path, coperture, qe, ai_premesse)
    elif tipo == "Liquidazione":
        return _gen_liquidazione(d, path, coperture, qe, ai_premesse)
    elif tipo == "Determina Dirigenziale":
        return _gen_determina_dirigenziale(d, path, coperture, qe, ai_premesse)
    else:
        # Fallback per tipi vecchi o generici
        if "Impegno" in str(tipo): return _gen_impegno_accertamento(d, path, coperture, qe, ai_premesse)
        if "Liquidazione" in str(tipo): return _gen_liquidazione(d, path, coperture, qe, ai_premesse)
        return _gen_generica(d, path, coperture, qe, ai_premesse)


def _titolo_det(d):
    t = "DETERMINAZIONE DIRIGENZIALE"
    if d.get("numero"): t += f" N. {d['numero']}"
    if d.get("anno"): t += f" / {d['anno']}"
    return t

def _gen_impegno_accertamento(d, path, coperture, qe, ai_premesse):
    b = DocBuilder(); _intestazione(b, d)
    oggetto = _o(d.get("affidamento_oggetto", d.get("oggetto")))
    importo = d.get("importo", 0) or 0
    cig = _o(d.get("cig"), "")
    rup = _o(d.get("rup"))

    b.center(_titolo_det(d), bold=True); b.empty()
    b.bold(f"OGGETTO: {oggetto.upper()} — Impegno di spesa e accertamento fondi. CIG: {cig}."); b.empty()
    b.center("IL DIRIGENTE", bold=True); b.empty()

    b.bold("PREMESSO CHE")
    _premesse_ai_o_manuali(b, d, ai_premesse,
        f"si rende necessario procedere all'impegno della spesa e al contestuale accertamento "
        f"delle risorse per l'intervento: {oggetto.lower()};")

    b.bold("RICHIAMATI")
    b.normal("– gli atti e i provvedimenti relativi alla programmazione dell'Ente;")
    b.normal(f"– la precedente determinazione dirigenziale (se presente) relativa all'affidamento in oggetto;")
    b.empty()

    _visti_sanremo(b)
    b.empty()

    _qe_tabella(b, qe)
    _coperture_tabella(b, coperture)

    b.bold("DATO ATTO CHE")
    b.normal(f"il Responsabile Unico del Progetto è {rup};")
    b.normal(f"la spesa complessiva ammonta a {fmtE(importo)};")
    b.normal("è stato esercitato il controllo preventivo di regolarità e correttezza amministrativa "
             "ai sensi dell'art. 147-bis del D.Lgs. 267/2000;")
    b.empty()

    _qe_tabella(b, qe)

    b.center("DETERMINA", bold=True); b.empty()
    b.mixed(('B', "1. "), ('', f"Di approvare le premesse che qui si intendono integralmente richiamate;")); b.empty()
    b.mixed(('B', "2. "), ('', f"Di impegnare la spesa di {fmtE(importo)} come da prospetto di copertura in calce;")); b.empty()
    b.mixed(('B', "3. "), ('', f"Di dare atto che il programma dei pagamenti è compatibile con gli stanziamenti di bilancio;")); b.empty()
    b.mixed(('B', "4. "), ('', "Di trasmettere la presente all'Ufficio Ragioneria per i visti di competenza."))

    _firma(b, d)
    
    _attestazione_copertura(b, d, coperture)
    
    return b.save(path)


# ═══════════════════════════════════════════════════════════════
# FUNZIONI DI SUPPORTO QE
# ═══════════════════════════════════════════════════════════════

def _qe_tabella(b, qe):
    if not qe: return
    
    headers = ["Sez.", "Descrizione", "Importo", "IVA %", "Sogg. IVA", "Totale"]
    rows = []
    
    # Ordina per sezione e ordine
    qe_sorted = sorted(qe, key=lambda x: (x.get("sezione", "A"), x.get("ordine", 0)))
    
    ta = 0; tb = 0; tiva = 0
    for v in qe_sorted:
        imp = v.get("importo", 0) or 0
        iva_p = v.get("aliquota_iva", 22) or 0
        sogg = "Sì" if v.get("soggetto_iva") else "No"
        
        iva_val = round(imp * iva_p / 100, 2) if v.get("soggetto_iva") else 0
        tot = imp + iva_val
        
        if v.get("sezione") == "A": ta += imp
        else: tb += imp
        tiva += iva_val
        
        rows.append([
            v.get("sezione", "A"),
            v.get("descrizione", ""),
            fmtE(imp),
            f"{iva_p}%",
            sogg,
            fmtE(tot)
        ])
    
    # Righe di totale
    rows.append(["", "", "", "", "", ""]) # Spazio
    rows.append(["", "TOTALE A) LAVORI E SERVIZI", fmtE(ta), "", "", ""])
    rows.append(["", "TOTALE B) SOMME A DISPOSIZIONE", fmtE(tb), "", "", ""])
    rows.append(["", "TOTALE IVA", fmtE(tiva), "", "", ""])
    rows.append(["", "TOTALE GENERALE QUADRO ECONOMICO", "", "", "", fmtE(ta + tb + tiva)])
    
    b.bold("QUADRO ECONOMICO")
    b.table(rows, headers=headers, widths=[1.5, 9, 2.5, 1.5, 1.5, 2.5])


def _coperture_tabella(b, coperture):
    if not coperture: return
    
    b.bold("COPERTURA FINANZIARIA")
    headers = ["Capitolo", "Descrizione", "Esercizio", "Importo"]
    rows = []
    tot = 0
    for c in coperture:
        imp = c.get("importo", 0) or 0
        tot += imp
        rows.append([
            str(c.get("capitolo", "___")),
            c.get("descrizione", ""),
            str(c.get("esercizio", "___")),
            fmtE(imp)
        ])
    rows.append(["", "TOTALE COPERTURA", "", fmtE(tot)])
    b.table(rows, headers=headers, widths=[3.5, 8.5, 2.5, 3.5])
    b.empty()


# ═══════════════════════════════════════════════════════════════
# DETERMINA A CONTRARRE (+ AFFIDAMENTO DIRETTO)
# ═══════════════════════════════════════════════════════════════

def _gen_determina_dirigenziale(d, path, coperture, qe, ai_premesse):
    b = DocBuilder(); _intestazione(b, d)

    oggetto = _o(d.get("affidamento_oggetto", d.get("oggetto")))
    rup = _o(d.get("rup")); qualifica_rup = _o(d.get("qualifica_rup"), "")
    rag_soc = _o(d.get("ragione_sociale")); piva = _o(d.get("partita_iva"), "")
    cf = _o(d.get("codice_fiscale"), "")
    sede = f"{_o(d.get('fornitore_citta'), '')} {_o(d.get('fornitore_indirizzo'), '')}".strip() or "___"
    legale_rappr = _o(d.get("legale_rappresentante"), "")
    cig = _o(d.get("cig"), ""); cup = _o(d.get("cup"), "")
    rif_norm = _o(d.get("rif_normativo", d.get("aff_rif_normativo")), "dell'art. 50 del D.Lgs. 36/2023")
    importo = d.get("importo", 0) or 0
    tipo_prest = _o(d.get("tipo_prestazione"), "")
    prot_prev = _o(d.get("prot_preventivo"), "___"); data_prev = fmtD(d.get("data_preventivo"))
    prot_durc = _o(d.get("prot_durc"), "___"); valid_durc = _o(d.get("validita_durc"), "___")
    forma = _o(d.get("forma_contratto"), "corrispondenza")
    tempi = _o(d.get("tempi_esecuzione"), "")
    prop_n = _o(d.get("id", "___")) # Proposta è l'ID nel nostro sistema

    cf_piva = ""
    if cf and cf != "___": cf_piva += f"C.F. {cf}"
    if piva and piva != "___": cf_piva += (" / " if cf_piva else "") + f"P.IVA {piva}"
    if not cf_piva: cf_piva = "C.F./P.IVA ___"

    is_lavori = "lavor" in tipo_prest.lower() if tipo_prest else False
    is_tecnici = "tecnic" in tipo_prest.lower() if tipo_prest else False
    soggetto_tipo = "la ditta" if is_lavori else ("il professionista" if is_tecnici else "l'operatore economico")
    opera = "i lavori" if is_lavori else ("l'incarico professionale" if is_tecnici else "il servizio/la fornitura")

    b.center(_titolo_det(d), bold=True); b.empty()
    b.bold(f"OGGETTO: {oggetto.upper()}" + (f" - CIG: {cig}" if cig and cig != "___" else "")); b.empty()
    b.normal(f"Proposta n. {prop_n}"); b.empty()
    b.center("IL DIRIGENTE", bold=True); b.empty()

    # PREMESSE / VISTI / RICHIAMATI
    if ai_premesse:
        _premesse_ai_o_manuali(b, d, ai_premesse, "")
    else:
        b.bold("PREMESSO CHE")
        b.normal(f"si rende necessario procedere all'affidamento di: {oggetto.lower()};"); b.empty()
        
        if rag_soc and rag_soc != "___":
            b.bold("DATO ATTO CHE")
            b.mixed(('', f"è stato richiesto preventivo a{soggetto_tipo[1:]} "), ('B', rag_soc),
                    ('', f", con sede in {sede}, {cf_piva}"),
                    ('', f", nella persona del legale rappresentante {legale_rappr};" if legale_rappr and legale_rappr != "___" else ";"))
            b.empty()
            b.normal(f"{soggetto_tipo.capitalize()} ha presentato il proprio preventivo "
                     f"(prot. n. {prot_prev} del {data_prev}) per un importo complessivo di {fmtE(importo)} IVA inclusa;")
            b.empty()

    b.bold("DATO ATTO che:")
    b.normal(f"- il soggetto responsabile del procedimento, ai sensi dell'art. 6 della legge 241/1990 è {rup}" +
             (f", {qualifica_rup}" if qualifica_rup and qualifica_rup != "___" else "") + ";")
    b.normal(f"- è stato esercitato, in merito alla presente determinazione, il controllo preventivo di regolarità tecnica "
             "e correttezza dell'azione amministrativa ai sensi dell'art. 147 bis, 1° comma, del D.Lgs. 267/2000;"); b.empty()

    _visti_sanremo(b); b.empty()

    if qe:
        _qe_tabella(b, qe)

    b.center("D E T E R M I N A", bold=True); b.empty()

    n = 1
    # Punto 1: Affidamento o Approvazione
    if rag_soc and rag_soc != "___":
        b.mixed(('B', f"{n}) "), ('', f"di affidare direttamente, ai sensi {rif_norm}, {opera} di cui in oggetto a{soggetto_tipo[1:]} "),
                ('B', rag_soc), ('', f", con sede in {sede}, {cf_piva}, per un importo complessivo di {fmtE(importo)} IVA inclusa;")); b.empty(); n += 1
    else:
        b.mixed(('B', f"{n}) "), ('', f"di approvare l'avvio del procedimento relativo a: {oggetto.lower()};")); b.empty(); n += 1

    # Punto 2: Copertura Finanziaria (se importo > 0)
    if importo > 0:
        b.mixed(('B', f"{n}) "), ('', "di impegnare la spesa complessiva di "), ('B', fmtE(importo)))
        if coperture:
            if len(coperture) > 1:
                headers = ["Miss/Prog/Tit", "Capitolo", "Bilancio", "Importo"]
                rows = [[f"{c.get('missione','_')}/{c.get('programma','_')}/{c.get('titolo','_')}", 
                         c.get('capitolo','_'), c.get('anno_bilancio','_'), fmtE(c.get('importo',0))] for c in coperture]
                b.table(rows, headers=headers, widths=[3, 7, 3, 3])
            else:
                b.normal(f"   {_fmt_cop(coperture[0])};")
        else:
            b.normal(f"   al Cap. {_o(d.get('capitolo_bilancio', d.get('aff_capitolo')))} del Bilancio {_o(d.get('esercizio'))};")
        b.empty(); n += 1

    # Punti Standard
    b.mixed(('B', f"{n}) "), ('', "di dare atto che si ottempererà alle disposizioni in materia di Amministrazione Trasparente;")); b.empty(); n += 1
    
    b.mixed(('B', f"{n}) "), ('', "di dare atto che è stato accertato che il programma dei pagamenti è compatibile con i relativi stanziamenti di bilancio;")); b.empty(); n += 1
    
    b.mixed(('B', f"{n}) "), ('', "di dare atto che la presente determinazione ha efficacia immediata ai sensi del vigente Regolamento sull'ordinamento degli uffici e dei servizi comunali."))

    _firma(b, d)
    
    if importo > 0:
        _attestazione_copertura(b, d, coperture)
        
    return b.save(path)


# ═══════════════════════════════════════════════════════════════
# AFFIDAMENTO (post-gara)
# ═══════════════════════════════════════════════════════════════

def _gen_affidamento(d, path, coperture, qe, ai_premesse):
    b = DocBuilder(); _intestazione(b, d)
    oggetto = _o(d.get("affidamento_oggetto", d.get("oggetto")))
    rag_soc = _o(d.get("ragione_sociale")); piva = _o(d.get("partita_iva"), "")
    sede = f"{_o(d.get('fornitore_citta'), '')} {_o(d.get('fornitore_indirizzo'), '')}".strip() or "___"
    cig = _o(d.get("cig"), ""); importo = d.get("importo", 0) or 0
    padre_n = _o(d.get("padre_numero"), "___"); padre_d = fmtD(d.get("padre_data"))
    rup = _o(d.get("rup")); qualifica_rup = _o(d.get("qualifica_rup"), "")

    b.center(_titolo_det(d), bold=True); b.empty()
    b.bold(f"OGGETTO: {oggetto.upper()} — Determinazione di affidamento. CIG: {cig}."); b.empty()
    b.center("IL DIRIGENTE", bold=True); b.empty()

    b.bold("RICHIAMATA")
    b.normal(f"la propria Determinazione Dirigenziale n. {padre_n} del {padre_d} con la quale "
             f"è stata avviata la procedura per l'affidamento di: {oggetto.lower()};"); b.empty()

    b.bold("PREMESSO CHE")
    _premesse_ai_o_manuali(b, d, ai_premesse,
        f"a seguito dell'espletamento della procedura di affidamento, l'operatore economico "
        f"{rag_soc}, con sede in {sede}, P.IVA {piva}, è risultato aggiudicatario;"); b.empty()

    b.bold("DATO ATTO CHE")
    b.normal(f"il Responsabile Unico del Progetto, ai sensi dell'art. 15 del D.Lgs. 36/2023, "
             f"è {rup}{', ' + qualifica_rup if qualifica_rup and qualifica_rup != '___' else ''};"); b.empty()

    dl = d.get("direttore_lavori")
    if dl and dl != "___":
        qdl = d.get("qualifica_dl", "")
        b.normal(f"il Direttore dei Lavori è {dl}{', ' + qdl if qdl else ''};"); b.empty()

    coll = d.get("collaboratori")
    if coll and coll != "___":
        b.normal(f"i collaboratori del RUP sono: {coll};"); b.empty()

    b.normal("sono state positivamente espletate le verifiche di cui agli artt. 94-98 "
             "del D.Lgs. 36/2023 in ordine ai requisiti di partecipazione;"); b.empty()
    _coperture_testo(b, d, coperture); b.empty()
    _visti_sanremo(b); b.empty()

    _qe_tabella(b, qe)

    b.center("DETERMINA", bold=True); b.empty()
    b.mixed(('B', "1. "), ('', f"Di aggiudicare definitivamente a {rag_soc}, con sede in {sede}, P.IVA {piva}, "
                               f"l'esecuzione di: {oggetto.lower()}, per l'importo complessivo di {fmtE(importo)} IVA inclusa;")); b.empty()
    b.mixed(('B', "2. "), ('', "Di dare atto che la spesa trova copertura come indicato in premessa;")); b.empty()
    b.mixed(('B', "3. "), ('', "Di trasmettere la presente all'Ufficio Ragioneria per il visto di regolarità contabile."))
    _firma(b, d); return b.save(path)


# ═══════════════════════════════════════════════════════════════
# LIQUIDAZIONE
# ═══════════════════════════════════════════════════════════════

def _gen_liquidazione(d, path, coperture, qe, ai_premesse):
    b = DocBuilder(); _intestazione(b, d)
    
    # Dati Fattura
    num_f = _o(d.get("numero_fattura"), "___")
    data_f = fmtD(d.get("data_fattura"))
    prot_f = _o(d.get("protocollo_pec"), _o(d.get("prot_fattura"), "___"))
    imp_netto = d.get("fat_netto", d.get("importo_netto", 0)) or 0
    aliq = d.get("fat_aliquota", d.get("aliquota_iva", 22)) or 22
    imp_iva = d.get("fat_iva", d.get("importo_iva", 0)) or 0
    imp_tot = d.get("fat_totale", d.get("importo_totale", 0)) or 0
    cassa = d.get("fat_cassa", 0) or 0
    
    # Dati Operatore
    rag_soc = _o(d.get("ragione_sociale"))
    piva = _o(d.get("partita_iva"), "")
    cf = _o(d.get("codice_fiscale"), "")
    sede = f"{_o(d.get('fornitore_citta', d.get('operatore_citta')), '')} " \
           f"{_o(d.get('fornitore_indirizzo', d.get('operatore_indirizzo')), '')}".strip() or "___"
    cf_piva = ""
    if cf and cf != "___": cf_piva += f"C.F. {cf}"
    if piva and piva != "___": cf_piva += (" / " if cf_piva else "") + f"P.IVA {piva}"
    if not cf_piva: cf_piva = "C.F./P.IVA ___"

    # Dati Affidamento
    oggetto = _o(d.get("affidamento_oggetto", d.get("oggetto")))
    tipo_prest = _o(d.get("tipo_prestazione"), "SERVIZI")
    cig = _o(d.get("cig"), "")
    rup = _o(d.get("rup"))
    padre_n = _o(d.get("padre_numero", d.get("aff_numero")), "___")
    padre_d = fmtD(d.get("padre_data", d.get("aff_data")))
    
    # Dati DURC e Regolarità
    data_cre = fmtD(d.get("data_cre", d.get("data_regolare_esecuzione")))
    durc_prot = _o(d.get("durc_prot"), "___________")
    durc_scad = fmtD(d.get("durc_scadenza"))

    # Contabile
    impegno = _o(d.get("impegno_spesa"), "___")
    esercizio = _o(d.get("esercizio"), "___")

    # Oggetto
    ogg_titolo = f"{tipo_prest.upper()} - {oggetto.upper()} - CIG {cig}: LIQUIDAZIONE FATTURA N. {num_f} DEL {data_f}"
    
    b.center(_titolo_det(d), bold=True); b.empty()
    b.bold(f"OGGETTO: {ogg_titolo}"); b.empty()
    b.center("IL DIRIGENTE", bold=True); b.empty()

    b.bold("RICHIAMATI i provvedimenti:")
    b.normal(f"– n. {padre_n} del {padre_d} con il quale è stato disposto l'affidamento di: {oggetto.lower()} "
             f"all'Impresa {rag_soc.upper()}, con sede in {sede} ({cf_piva}), impegnando la relativa spesa;")
    if d.get("altri_provvedimenti"):
        b.normal(f"– {d['altri_provvedimenti']};")
    b.empty()

    b.bold("RILEVATO")
    b.normal("che le clausole contrattuali prevedono la liquidazione dei corrispettivi a seguito della verifica "
             "della regolare esecuzione delle prestazioni;"); b.empty()

    b.bold("DATO ATTO")
    b.normal(f"che in data {data_cre} il Responsabile Unico del Progetto, {rup}, ha attestato la regolare "
             f"esecuzione delle prestazioni in argomento per la fase di riferimento;"); b.empty()

    b.bold("VISTE:")
    b.normal(f"– la fattura n. {num_f} del {data_f}, acquisita in atti con prot. n. {prot_f}, per un importo di "
             f"{fmtE(imp_netto)} oltre IVA {int(aliq)}% pari a {fmtE(imp_iva)}" + 
             (f" e Cassa pari a {fmtE(cassa)}" if cassa else "") + 
             f", per un ammontare complessivo pari a {fmtE(imp_tot)};"); b.empty()

    b.bold("DATO ATTO CHE:")
    b.normal(f"– il Responsabile Unico del Progetto, {rup}, ha emesso in data {data_cre} l'attestazione di regolare esecuzione;")
    b.normal(f"– il CIG relativo al contratto è {cig};")
    b.normal(f"– è stata verificata la Regolarità Contributiva come risulta dal DURC n. {durc_prot} con scadenza {durc_scad};")
    b.normal("– si è provveduto ad ottemperare alle disposizioni del D.Lgs. 33/2013 in relazione agli atti richiamati;")
    b.normal("– è stato esercitato il controllo preventivo di regolarità e correttezza amministrativa ai sensi dell'art. 147-bis del D.Lgs. 267/2000;"); b.empty()

    b.bold("DATO ATTO")
    b.normal("che è stato accertato, ai sensi dell'art. 9 del D.L. 78/2009, che il programma dei pagamenti "
             "è compatibile con i relativi stanziamenti di Bilancio e con le regole di finanza pubblica;"); b.empty()

    b.bold("RITENUTO")
    b.normal(f"pertanto di dover procedere alla liquidazione della fattura n. {num_f} del {data_f} "
             f"per complessivi {fmtE(imp_tot)} a favore di {rag_soc};"); b.empty()

    _visti_sanremo(b); b.empty()

    b.center("DETERMINA", bold=True); b.empty()
    
    b.mixed(('B', "1. "), ('', f"Di procedere alla liquidazione della fattura n. {num_f} del {data_f}, acquisita al prot. n. {prot_f}, "
             f"emessa dalla ditta {rag_soc} ({cf_piva}) per un importo di complessivi {fmtE(imp_tot)}, "
             f"dando atto che la spesa è stata finanziata con determinazione n. {padre_n} del {padre_d}, "
             f"con impegno di spesa n. {impegno}, esercizio {esercizio};")); b.empty()
    
    b.mixed(('B', "2. "), ('', "Di incaricare il Servizio Spesa di procedere all'emissione del mandato di pagamento "
             "sul conto corrente bancario dedicato, ai sensi della normativa vigente in materia di tracciabilità dei flussi finanziari;")); b.empty()
    
    b.mixed(('B', "3. "), ('', "Di ottemperare alle disposizioni previste dal D.Lgs. 33/2013 'Amministrazione Trasparente';")); b.empty()
    
    b.mixed(('B', "4. "), ('', "Di dare atto che il programma dei pagamenti è compatibile con i relativi stanziamenti di Bilancio;")); b.empty()
    
    b.mixed(('B', "5. "), ('', "Di dare atto che la presente determinazione ha efficacia immediata ai sensi del Regolamento comunale vigente."))

    _firma(b, d); return b.save(path)


# ═══════════════════════════════════════════════════════════════
# INTEGRAZIONE, REVOCA, GENERICA
# ═══════════════════════════════════════════════════════════════


def _gen_integrazione(d, path, coperture, qe, ai_premesse):
    b = DocBuilder(); _intestazione(b, d)
    oggetto = _o(d.get("affidamento_oggetto", d.get("oggetto"))); importo = d.get("importo", 0) or 0
    padre_n = _o(d.get("padre_numero"), "___"); padre_d = fmtD(d.get("padre_data"))
    b.center(_titolo_det(d), bold=True); b.empty()
    b.bold(f"OGGETTO: Integrazione impegno di spesa — {oggetto}"); b.empty()
    b.center("IL DIRIGENTE", bold=True); b.empty()
    b.bold("RICHIAMATA"); b.normal(f"la Determinazione Dirigenziale n. {padre_n} del {padre_d};"); b.empty()
    b.bold("PREMESSO CHE")
    _premesse_ai_o_manuali(b, d, ai_premesse, f"si rende necessario integrare l'impegno di spesa relativo a: {oggetto.lower()};")
    b.bold("DATO ATTO CHE"); _coperture_testo(b, d, coperture); b.empty()
    _visti_sanremo(b, include_tracciabilita=False); b.empty()
    
    _qe_tabella(b, qe)
    
    b.center("DETERMINA", bold=True); b.empty()
    b.mixed(('B', "1. "), ('', f"Di integrare l'impegno di spesa per {fmtE(importo)};")); b.empty()
    b.mixed(('B', "2. "), ('', "Di trasmettere la presente all'Ufficio Ragioneria per il visto di regolarità contabile."))
    _firma(b, d); return b.save(path)


def _gen_revoca(d, path, coperture, qe, ai_premesse):
    b = DocBuilder(); _intestazione(b, d)
    oggetto = _o(d.get("affidamento_oggetto", d.get("oggetto")))
    padre_n = _o(d.get("padre_numero"), "___"); padre_d = fmtD(d.get("padre_data"))
    b.center(_titolo_det(d), bold=True); b.empty()
    b.bold(f"OGGETTO: Revoca — {oggetto}"); b.empty()
    b.center("IL DIRIGENTE", bold=True); b.empty()
    b.bold("RICHIAMATA"); b.normal(f"la Determinazione Dirigenziale n. {padre_n} del {padre_d};"); b.empty()
    b.bold("PREMESSO CHE")
    _premesse_ai_o_manuali(b, d, ai_premesse, f"si rende necessario procedere alla revoca del provvedimento relativo a: {oggetto.lower()};")
    _visti_sanremo(b, include_tracciabilita=False); b.empty()
    
    _qe_tabella(b, qe)
    
    b.center("DETERMINA", bold=True); b.empty()
    b.mixed(('B', "1. "), ('', f"Di revocare la Determinazione Dirigenziale n. {padre_n} del {padre_d} relativa a: {oggetto.lower()};")); b.empty()
    b.mixed(('B', "2. "), ('', "Di trasmettere la presente all'Ufficio Ragioneria per i conseguenti adempimenti contabili."))
    _firma(b, d); return b.save(path)


def _gen_generica(d, path, coperture, qe, ai_premesse):
    b = DocBuilder(); _intestazione(b, d)
    tipo = d.get("tipo_determina", "Determina"); oggetto = _o(d.get("affidamento_oggetto", d.get("oggetto")))
    b.center(_titolo_det(d), bold=True); b.empty()
    b.bold(f"OGGETTO: {tipo} — {oggetto}"); b.empty()
    b.center("IL DIRIGENTE", bold=True); b.empty()
    b.bold("PREMESSO CHE")
    _premesse_ai_o_manuali(b, d, ai_premesse, f"si rende necessario adottare il presente provvedimento relativo a: {oggetto.lower()};")
    _visti_sanremo(b); b.empty()
    
    _qe_tabella(b, qe)
    
    b.center("DETERMINA", bold=True); b.empty()
    b.normal("[Completare il dispositivo]")
    _firma(b, d); return b.save(path)


# ═══════════════════════════════════════════════════════════════
# VERBALI DL
# ═══════════════════════════════════════════════════════════════

def genera_verbale(vdata, adata, output_path):
    b = DocBuilder()
    b.center("COMUNE DI SANREMO", bold=True); b.center("Provincia di Imperia"); b.empty()
    b.bold(_o(adata.get("qualifica_dirigente"), "Settore Sviluppo Economico, Ambientale e Floricoltura"))
    b.bold("Servizio Demanio Marittimo"); b.empty()

    tipo = vdata.get("tipo_verbale", ""); data_v = fmtD(vdata.get("data_verbale", ""))
    num = vdata.get("numero_progressivo", 1)
    redattore = _o(vdata.get("redattore"), _o(adata.get("rup"), "___"))
    oggetto = _o(adata.get("oggetto", adata.get("affidamento_oggetto")))
    cig = _o(adata.get("cig"), "")
    rag_soc = _o(adata.get("ragione_sociale", adata.get("operatore_nome")))
    sede = f"{_o(adata.get('operatore_citta', adata.get('fornitore_citta')), '')} " \
           f"{_o(adata.get('operatore_indirizzo', adata.get('fornitore_indirizzo')), '')}".strip() or "___"
    piva = _o(adata.get("operatore_piva", adata.get("partita_iva")), "")
    dirigente = _o(adata.get("dirigente"), "Il Dirigente")

    b.center(f"VERBALE DI {tipo.upper()}", bold=True)
    if tipo == "SAL": b.center(f"N. {num}", bold=True)
    b.empty()
    b.normal(f"L'anno _______ il giorno _______ del mese di _______ ({data_v})"); b.empty()
    b.bold("PRESENTI:")
    b.normal(f"– il Direttore dei Lavori, {redattore};")
    b.normal(f"– per l'Impresa {rag_soc}, con sede in {sede}, P.IVA {piva}: il legale rappresentante o suo delegato;")
    b.empty()

    b.bold("PREMESSO CHE")
    b.normal(f"con Determinazione Dirigenziale è stato affidato all'Impresa {rag_soc} "
             f"{('(CIG ' + cig + ') ') if cig and cig != '___' else ''}l'esecuzione dei {oggetto.lower()};"); b.empty()

    if tipo == "Consegna dei luoghi":
        b.bold("SI PROCEDE")
        b.normal("alla consegna dei luoghi e alla delimitazione delle aree interessate dall'intervento, come da elaborati progettuali approvati."); b.empty()
        b.normal("L'Appaltatore dichiara di aver preso visione dei luoghi e di accettarli nello stato in cui si trovano, e di averli trovati idonei per dare inizio ai lavori."); b.empty()
        b.normal("Il Direttore dei Lavori impartisce le disposizioni necessarie per l'inizio dei lavori e prescrive le cautele da adottare.")
    elif tipo == "SAL":
        imp_sal = vdata.get("importo_sal", 0) or 0
        b.bold("IL DIRETTORE DEI LAVORI CERTIFICA")
        b.normal(f"che i lavori eseguiti alla data odierna ammontano, al netto del ribasso d'asta e degli oneri della sicurezza, a {fmtE(imp_sal)}."); b.empty()
        b.normal("Le lavorazioni eseguite risultano conformi al contratto e al capitolato d'appalto.")
    elif tipo == "Sospensione lavori":
        b.bold("IL DIRETTORE DEI LAVORI ORDINA")
        b.normal("la sospensione dei lavori per le seguenti motivazioni:"); b.empty()
        b.normal(_o(vdata.get("note"), "[Indicare le motivazioni della sospensione]")); b.empty()
        b.normal("L'Appaltatore prende atto della sospensione e conferma di ritenere adeguatamente custoditi i materiali e le attrezzature presenti in cantiere.")
    elif tipo == "Ripresa lavori":
        b.bold("IL DIRETTORE DEI LAVORI")
        b.normal("essendo venute meno le cause di sospensione, ordina la ripresa dei lavori."); b.empty()
        b.normal("L'Appaltatore conferma di poter riprendere immediatamente le lavorazioni e dichiara che il cantiere si trova nelle condizioni idonee alla ripresa.")
    elif tipo == "Fine lavori":
        b.bold("IL DIRETTORE DEI LAVORI ATTESTA")
        b.normal("che i lavori in oggetto sono stati ultimati in data odierna."); b.empty()
        b.normal("I lavori risultano eseguiti conformemente al contratto, al capitolato d'appalto e agli eventuali atti aggiuntivi."); b.empty()
        b.normal("L'Appaltatore chiede che si proceda alla verifica di conformità / certificato di regolare esecuzione delle opere realizzate.")
    elif tipo == "CRE":
        b.bold("IL DIRETTORE DEI LAVORI")
        b.normal("a seguito di accurato sopralluogo e verifica delle opere realizzate, attesta la regolare esecuzione dei lavori e la conformità degli stessi "
                 "al contratto, al capitolato d'appalto e agli elaborati progettuali approvati."); b.empty()
        b.normal("Dichiara pertanto che le prestazioni sono state eseguite a regola d'arte, nel rispetto delle norme vigenti in materia e delle prescrizioni impartite in corso d'opera."); b.empty()
        b.normal("Non si rilevano vizi, difetti o difformità rispetto a quanto contrattualmente previsto.")

    b.empty()
    if vdata.get("note") and tipo not in ("Sospensione lavori",):
        b.bold("NOTE:"); b.normal(vdata["note"]); b.empty()

    b.normal("Letto, confermato e sottoscritto."); b.empty(); b.empty()
    b.normal(f"Il Direttore dei Lavori"); b.normal(f"{redattore}"); b.empty()
    b.normal(f"Per l'Impresa"); b.normal(f"{rag_soc}"); b.empty(); b.empty()
    b.right(f"Il Dirigente"); b.right(dirigente)
    return b.save(output_path)
